from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from fastapi.testclient import TestClient
from openpyxl import Workbook
import pytest

from app.assistant.orchestrator import AssistantOrchestrator
from app.main import app


client = TestClient(app)


def test_pdf_content_extraction_finds_query(monkeypatch, tmp_path) -> None:
    from app.tools.files.content_extractors import extract_text_from_pdf

    pdf_path = tmp_path / "vertrag.pdf"
    pdf_path.write_bytes(b"%PDF-1.4")

    class FakePage:
        def extract_text(self) -> str:
            return "Kaufvertrag Hauskauf"

    class FakeReader:
        pages = [FakePage()]

        def __init__(self, _path: Path) -> None:
            pass

    monkeypatch.setattr("app.tools.files.content_extractors.PdfReader", FakeReader)

    assert "Kaufvertrag" in extract_text_from_pdf(pdf_path)["text"]


def test_docx_content_extraction_finds_query(tmp_path) -> None:
    from app.tools.files.content_extractors import extract_text_from_docx

    path = tmp_path / "vertrag.docx"
    document = Document()
    document.add_paragraph("Kaufvertrag Hauskauf")
    document.save(path)

    assert "Kaufvertrag" in extract_text_from_docx(path)


def test_xlsx_content_extraction_finds_query(tmp_path) -> None:
    from app.tools.files.content_extractors import extract_text_from_xlsx

    path = tmp_path / "daten.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet["A1"] = "Kaufvertrag"
    workbook.save(path)

    assert "Kaufvertrag" in extract_text_from_xlsx(path)


def test_txt_content_extraction_finds_query(tmp_path) -> None:
    from app.tools.files.content_extractors import extract_text_from_text_file

    path = tmp_path / "notiz.txt"
    path.write_text("Kaufvertrag Hauskauf", encoding="utf-8")

    assert "Kaufvertrag" in extract_text_from_text_file(path)


def test_content_search_only_searches_allowed_dirs(monkeypatch, tmp_path) -> None:
    from app.tools.files.content_search_tool import ContentSearchTool

    allowed = tmp_path / "allowed"
    outside = tmp_path / "outside"
    allowed.mkdir()
    outside.mkdir()
    (outside / "vertrag.txt").write_text("Kaufvertrag", encoding="utf-8")
    monkeypatch.setenv("FILE_SEARCH_ALLOWED_DIRS", str(allowed))

    result = ContentSearchTool().search_file_contents("Kaufvertrag")

    assert result["count"] == 0


def test_content_search_path_traversal_blocked(monkeypatch, tmp_path) -> None:
    from app.tools.files.content_search_tool import ContentSearchTool

    allowed = tmp_path / "allowed"
    allowed.mkdir()
    monkeypatch.setenv("FILE_SEARCH_ALLOWED_DIRS", str(allowed))

    result = ContentSearchTool().inspect_file("..\\secret.txt", query="Kaufvertrag")

    assert result["blocked"] is True


def test_file_too_large_skipped(monkeypatch, tmp_path) -> None:
    from app.tools.files.content_extractors import extract_text

    path = tmp_path / "gross.txt"
    path.write_text("Kaufvertrag", encoding="utf-8")
    monkeypatch.setenv("FILE_CONTENT_MAX_FILE_SIZE_MB", "0")

    result = extract_text(path)

    assert result["skipped"] is True
    assert "zu gross" in result["reason"]


def test_extraction_error_does_not_crash(monkeypatch, tmp_path) -> None:
    from app.tools.files.content_extractors import extract_text

    path = tmp_path / "kaputt.txt"
    path.write_text("Kaufvertrag", encoding="utf-8")
    monkeypatch.setattr("app.tools.files.content_extractors.extract_text_from_text_file", lambda _path: (_ for _ in ()).throw(RuntimeError("kaputt")))

    result = extract_text(path)

    assert result["error"] is True


def test_content_extraction_cache_reuses_unchanged_file(monkeypatch, tmp_path) -> None:
    from app.tools.files import content_extractors

    path = tmp_path / "cached.txt"
    path.write_text("Kaufvertrag", encoding="utf-8")
    calls = {"count": 0}
    monkeypatch.setenv("FILE_CONTENT_CACHE_ENABLED", "true")
    monkeypatch.setenv("FILE_CONTENT_CACHE_MAX_ITEMS", "200")
    content_extractors.clear_content_cache()

    def fake_extract(_path: Path) -> str:
        calls["count"] += 1
        return "Kaufvertrag"

    monkeypatch.setattr("app.tools.files.content_extractors.extract_text_from_text_file", fake_extract)

    first = content_extractors.extract_text(path)
    second = content_extractors.extract_text(path)

    assert first["text"] == second["text"]
    assert calls["count"] == 1
    assert second["cache"] == "hit"


def test_content_search_endpoint_returns_200(monkeypatch, tmp_path) -> None:
    allowed = tmp_path / "allowed"
    allowed.mkdir()
    (allowed / "vertrag.txt").write_text("Kaufvertrag", encoding="utf-8")
    monkeypatch.setenv("FILE_SEARCH_ALLOWED_DIRS", str(allowed))

    response = client.get("/assistant/files/content-search", params={"q": "Kaufvertrag", "extension": ".txt"})

    assert response.status_code == 200
    assert response.json()["count"] == 1


def test_content_search_route_for_pdf_query(monkeypatch, tmp_path) -> None:
    allowed = tmp_path / "allowed"
    allowed.mkdir()
    (allowed / "energieausweis.txt").write_text("Energieausweis", encoding="utf-8")
    monkeypatch.setenv("FILE_SEARCH_ALLOWED_DIRS", str(allowed))

    def fake_search(self: Any, query: str, extensions: list[str] | None = None, limit: int = 25) -> dict[str, Any]:
        return {
            "query": query,
            "extensions": extensions,
            "count": 1,
            "files": [{"name": "haus.pdf", "path": str(allowed / "haus.pdf"), "extension": ".pdf", "score": 50, "snippets": ["Energieausweis"]}],
            "searched_dirs": [str(allowed)],
            "skipped": [],
            "message": "1 Dateien mit Inhaltstreffern gefunden.",
        }

    monkeypatch.setattr("app.tools.files.content_search_tool.ContentSearchTool.search_file_contents", fake_search)

    result = AssistantOrchestrator().handle_message("Jarvis, suche in PDFs nach Kaufvertrag.")

    assert result["tool"] == "file_content_search"
    assert result["result"]["query"] == "Kaufvertrag"
    assert result["result"]["extensions"] == [".pdf"]


def test_path_only_match_ranked_lower_than_filename_match(monkeypatch, tmp_path) -> None:
    from app.tools.files.file_search_tool import FileSearchTool

    base = tmp_path / "Hauskauf"
    base.mkdir()
    (base / "allgemein.pdf").write_text("x", encoding="utf-8")
    (base / "Hauskauf Vertrag.pdf").write_text("x", encoding="utf-8")
    monkeypatch.setenv("FILE_SEARCH_ALLOWED_DIRS", str(base))

    result = FileSearchTool().search_files("Hauskauf", extensions=[".pdf"])

    assert result["files"][0]["name"] == "Hauskauf Vertrag.pdf"
    assert "filename" in result["files"][0]["match_sources"]
    assert result["files"][1]["path_match_only"] is True


def test_invalid_pdf_header_is_skipped(tmp_path) -> None:
    from app.tools.files.content_extractors import extract_text

    path = tmp_path / "kaputt.pdf"
    path.write_bytes(b"\x00\x05\x16\x07\x00")

    result = extract_text(path)

    assert result["success"] is False
    assert result["error"] is True
    assert result["reason"] == "invalid_pdf_header"
    assert result["text"] == ""


def test_invalid_pdf_does_not_raise_exception(tmp_path) -> None:
    from app.tools.files.content_extractors import extract_text_from_pdf

    path = tmp_path / "kaputt.pdf"
    path.write_bytes(b"not a pdf")

    result = extract_text_from_pdf(path)

    assert result["success"] is False
    assert result["reason"] == "invalid_pdf_header"


def test_invalid_pdf_appears_in_skipped_list(monkeypatch, tmp_path) -> None:
    from app.tools.files.content_search_tool import ContentSearchTool

    allowed = tmp_path / "allowed"
    allowed.mkdir()
    bad_pdf = allowed / "kaputt.pdf"
    bad_pdf.write_bytes(b"not a pdf")
    monkeypatch.setenv("FILE_SEARCH_ALLOWED_DIRS", str(allowed))

    result = ContentSearchTool().search_file_contents("Kaufvertrag", extensions=[".pdf"])

    assert result["skipped_count"] == 1
    assert result["skipped"][0]["path"] == str(bad_pdf)
    assert result["skipped"][0]["reason"] == "invalid_pdf_header"


def test_content_search_continues_after_invalid_pdf(monkeypatch, tmp_path) -> None:
    from app.tools.files.content_search_tool import ContentSearchTool

    allowed = tmp_path / "allowed"
    allowed.mkdir()
    (allowed / "kaputt.pdf").write_bytes(b"not a pdf")
    (allowed / "vertrag.txt").write_text("Kaufvertrag Hauskauf", encoding="utf-8")
    monkeypatch.setenv("FILE_SEARCH_ALLOWED_DIRS", str(allowed))

    result = ContentSearchTool().search_file_contents("Kaufvertrag")

    assert result["count"] == 1
    assert result["skipped_count"] == 1
    assert result["files"][0]["name"] == "vertrag.txt"


def test_all_candidate_pdfs_fail_message_mentions_placeholders(monkeypatch, tmp_path) -> None:
    from app.tools.files.content_search_tool import ContentSearchTool

    allowed = tmp_path / "allowed"
    allowed.mkdir()
    (allowed / "placeholder.pdf").write_bytes(b"")
    monkeypatch.setenv("FILE_SEARCH_ALLOWED_DIRS", str(allowed))

    result = ContentSearchTool().search_file_contents("Kaufvertrag", extensions=[".pdf"])

    assert "OneDrive-Platzhalter" in result["message"]
    assert result["skipped"][0]["reason"] == "empty_or_placeholder_file"


def test_textless_valid_pdf_requires_ocr(tmp_path) -> None:
    from pypdf import PdfWriter

    from app.tools.files.content_extractors import extract_text

    path = tmp_path / "scan.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with path.open("wb") as file:
        writer.write(file)

    result = extract_text(path)

    assert result["skipped"] is True
    assert result["reason"] == "ocr_required"
    assert result["message"] == "Das PDF enthält keinen extrahierbaren Text. OCR wird noch nicht unterstützt."
    assert result["page_count"] == 1


def test_docx_extracts_headings_paragraphs_and_tables(tmp_path) -> None:
    from app.tools.files.content_extractors import extract_text

    path = tmp_path / "vertrag.docx"
    document = Document()
    document.add_heading("Kaufvertrag", level=1)
    document.add_paragraph("Der Notartermin ist am Montag.")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Kategorie"
    table.rows[0].cells[1].text = "Wert"
    table.rows[1].cells[0].text = "Preis"
    table.rows[1].cells[1].text = "250000"
    document.save(path)

    result = extract_text(path)

    assert "[Überschrift: Kaufvertrag]" in result["text"]
    assert "Der Notartermin ist am Montag." in result["text"]
    assert "Zeile 1 | Kategorie: Preis | Wert: 250000" in result["text"]


def test_xlsx_extracts_structured_rows_from_multiple_sheets(tmp_path) -> None:
    from app.tools.files.content_extractors import extract_text

    path = tmp_path / "kosten.xlsx"
    workbook = Workbook()
    costs = workbook.active
    costs.title = "Kosten"
    costs.append(["Kategorie", "Menge", "Preis", "Datum", "Aktiv"])
    costs.append(["Boden", 42, 28.90, date(2026, 6, 21), True])
    notes = workbook.create_sheet("Notizen")
    notes.append(["Thema", "Text"])
    notes.append(["Haus", "Energieausweis prüfen"])
    workbook.save(path)

    result = extract_text(path)

    assert result["sheet_count"] == 2
    assert result["row_count"] == 2
    assert "[Arbeitsblatt: Kosten]" in result["text"]
    assert "Zeile 1 | Kategorie: Boden | Menge: 42 | Preis: 28.90 | Datum: 2026-06-21 | Aktiv: true" in result["text"]
    assert "[Arbeitsblatt: Notizen]" in result["text"]


def test_xlsm_is_read_only_and_workbook_closes(monkeypatch, tmp_path) -> None:
    from app.tools.files.content_extractors import extract_text_from_xlsx

    path = tmp_path / "daten.xlsm"
    path.write_bytes(b"placeholder")
    calls: dict[str, Any] = {}

    class FakeSheet:
        title = "Daten"

        def iter_rows(self, values_only: bool = True):
            assert values_only is True
            return iter([("Name", "Wert"), ("PV", 42)])

    class FakeWorkbook:
        worksheets = [FakeSheet()]

        def close(self) -> None:
            calls["closed"] = True

    def fake_load_workbook(*args: Any, **kwargs: Any) -> FakeWorkbook:
        calls["args"] = args
        calls["kwargs"] = kwargs
        return FakeWorkbook()

    monkeypatch.setattr("app.tools.files.content_extractors.load_workbook", fake_load_workbook)

    text = extract_text_from_xlsx(path)

    assert "[Arbeitsblatt: Daten]" in text
    assert calls["kwargs"] == {"read_only": True, "data_only": True, "keep_vba": False}
    assert calls["closed"] is True


def test_xlsx_workbook_closes_when_iteration_fails(monkeypatch, tmp_path) -> None:
    from app.tools.files.content_extractors import extract_text_from_xlsx

    path = tmp_path / "fehler.xlsx"
    path.write_bytes(b"placeholder")
    calls: dict[str, Any] = {}

    class BrokenWorkbook:
        @property
        def worksheets(self):
            raise ValueError("kaputt")

        def close(self) -> None:
            calls["closed"] = True

    monkeypatch.setattr("app.tools.files.content_extractors.load_workbook", lambda *args, **kwargs: BrokenWorkbook())

    with pytest.raises(ValueError, match="kaputt"):
        extract_text_from_xlsx(path)

    assert calls["closed"] is True


def test_csv_supports_comma_semicolon_and_legacy_encoding(tmp_path) -> None:
    from app.tools.files.content_extractors import extract_text

    comma = tmp_path / "comma.csv"
    comma.write_text("Kategorie,Menge\nBoden,42\n", encoding="utf-8")
    semicolon = tmp_path / "semicolon.csv"
    semicolon.write_text("Kategorie;Menge\nFenster;8\n", encoding="utf-8")
    legacy = tmp_path / "legacy.csv"
    legacy.write_bytes("Kategorie;Hinweis\nBöden;prüfen\n".encode("cp1252"))

    assert "Zeile 1 | Kategorie: Boden | Menge: 42" in extract_text(comma)["text"]
    assert "Zeile 1 | Kategorie: Fenster | Menge: 8" in extract_text(semicolon)["text"]
    assert "Böden" in extract_text(legacy)["text"]


def test_txt_markdown_and_json_keep_readable_content(tmp_path) -> None:
    from app.tools.files.content_extractors import extract_text

    text_path = tmp_path / "notiz.txt"
    markdown_path = tmp_path / "notiz.md"
    json_path = tmp_path / "daten.json"
    text_path.write_text("Erster Absatz.\n\nZweiter Absatz.", encoding="utf-8")
    markdown_path.write_text("# Haus\n\nEnergieausweis prüfen.", encoding="utf-8")
    json_path.write_text('{"haus": {"status": "prüfen", "zimmer": 4}}', encoding="utf-8")

    assert "Erster Absatz.\n\nZweiter Absatz." in extract_text(text_path)["text"]
    assert "# Haus" in extract_text(markdown_path)["text"]
    assert '"status": "prüfen"' in extract_text(json_path)["text"]
