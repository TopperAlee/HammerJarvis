import csv
import io
import json
import os
import re
from datetime import date, datetime, time
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader
from pypdf.errors import PdfReadError, PdfStreamError

from app.assistant.performance.timing import time_operation

SUPPORTED_CONTENT_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xlsm", ".csv", ".txt", ".md", ".json"}
TEXT_PREVIEW_CHARS = int(os.getenv("FILE_CONTENT_PREVIEW_CHARS", "4000"))
_CONTENT_CACHE: dict[tuple[str, int, int], dict[str, Any]] = {}
_MAX_EXTRACTED_TEXT_CHARS = TEXT_PREVIEW_CHARS * 20


def extract_text_from_pdf(path: Path) -> dict[str, Any]:
    """Extract PDF pages in order without emitting parser errors to stdout."""

    try:
        if path.stat().st_size == 0:
            return _pdf_error("empty_or_placeholder_file", "Datei ist leer oder ein Platzhalter.")
        with path.open("rb") as file:
            if file.read(4) != b"%PDF":
                return _pdf_error("invalid_pdf_header", "Datei hat keine gültige PDF-Signatur.")

        reader = PdfReader(path)
        page_texts = [page.extract_text() or "" for page in reader.pages]
        text = "\n\n".join(page_text.strip() for page_text in page_texts if page_text.strip())
        if not text:
            return {
                "success": False,
                "text": "",
                "error": False,
                "skipped": True,
                "reason": "ocr_required",
                "message": "Das PDF enthält keinen extrahierbaren Text. OCR wird noch nicht unterstützt.",
                "page_count": len(page_texts),
                "extraction_status": "ocr_required",
            }
        return {
            "success": True,
            "text": text,
            "error": False,
            "skipped": False,
            "page_count": len(page_texts),
            "extraction_status": "extracted",
        }
    except (PdfReadError, PdfStreamError, EOFError, ValueError):
        return _pdf_error("pdf_parse_error", "PDF konnte nicht gelesen werden.")
    except Exception:
        return _pdf_error("pdf_parse_error", "PDF konnte nicht gelesen werden.")


def extract_text_from_docx(path: Path) -> str:
    """Return paragraphs and tables, but never inspect embedded DOCX package files."""

    document = Document(path)
    lines: list[str] = []
    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if not value:
            continue
        if _is_heading(paragraph.style.name if paragraph.style else ""):
            lines.append(f"[Überschrift: {value}]")
        else:
            lines.append(value)

    for table in document.tables:
        rows = [[_cell_text(cell.text) for cell in row.cells] for row in table.rows]
        nonblank_rows = [row for row in rows if any(cell for cell in row)]
        if not nonblank_rows:
            continue
        headers = _headers_from_row(nonblank_rows[0])
        for row_number, row in enumerate(nonblank_rows[1:], start=1):
            lines.append(_structured_row(row_number, headers, row))
    return "\n".join(lines)


def extract_text_from_xlsx(path: Path) -> str:
    """Extract cell values read-only; formulas are never evaluated or macros executed."""

    workbook = None
    try:
        workbook = load_workbook(path, read_only=True, data_only=True, keep_vba=False)
        text, _metadata = _spreadsheet_sections(workbook)
        return text
    finally:
        if workbook is not None:
            workbook.close()


def extract_text_from_text_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".csv":
        return _extract_csv(path)
    if suffix == ".json":
        with path.open("r", encoding="utf-8", errors="replace") as file:
            data = json.load(file)
        return _bounded_json(data)
    return _read_text_tolerant(path)


def extract_text(path: Path) -> dict[str, Any]:
    """Return a compatible, structured extraction result without exposing errors verbatim."""

    suffix = path.suffix.lower()
    base = {"path": str(path), "extension": suffix}
    try:
        if suffix not in SUPPORTED_CONTENT_EXTENSIONS:
            return {
                **base,
                "success": False,
                "supported": False,
                "skipped": True,
                "error": False,
                "reason": "nicht unterstuetzter Dateityp",
                "extraction_status": "skipped",
                "text": "",
            }
        max_bytes = int(float(os.getenv("FILE_CONTENT_MAX_FILE_SIZE_MB", "25")) * 1024 * 1024)
        stat = path.stat()
        if stat.st_size > max_bytes:
            return {
                **base,
                "success": False,
                "supported": True,
                "skipped": True,
                "error": False,
                "reason": "Datei zu gross",
                "extraction_status": "skipped",
                "text": "",
            }
        cache_key = (str(path.resolve()), stat.st_mtime_ns, stat.st_size)
        cached = _get_cached(cache_key)
        if cached is not None:
            return {**cached, "cache": "hit"}

        metadata: dict[str, Any] = {}
        with time_operation(f"content_extract.{suffix}", "content_search"):
            if suffix == ".pdf":
                pdf_result = extract_text_from_pdf(path)
                metadata = {key: value for key, value in pdf_result.items() if key != "text"}
                if not pdf_result.get("success"):
                    return {
                        **base,
                        "success": False,
                        "supported": True,
                        "skipped": bool(pdf_result.get("skipped", True)),
                        "error": bool(pdf_result.get("error", True)),
                        "reason": pdf_result.get("reason"),
                        "message": pdf_result.get("message"),
                        "text": "",
                        "page_count": pdf_result.get("page_count", 0),
                        "extraction_status": pdf_result.get("extraction_status", "error"),
                    }
                text = str(pdf_result["text"])
            elif suffix == ".docx":
                text = extract_text_from_docx(path)
            elif suffix in {".xlsx", ".xlsm"}:
                text, metadata = _extract_spreadsheet_with_metadata(path)
            elif suffix == ".csv":
                text, metadata = _extract_csv_with_metadata(path)
            else:
                text = extract_text_from_text_file(path)

        text = clean_extracted_text(text)
        result = {
            **base,
            "success": True,
            "supported": True,
            "skipped": False,
            "error": False,
            "extraction_status": "extracted",
            "text": text,
            "preview": text[:TEXT_PREVIEW_CHARS],
            **metadata,
        }
        _set_cached(cache_key, result)
        return {**result, "cache": "miss"} if _content_cache_enabled() else result
    except Exception:
        return {
            **base,
            "success": False,
            "supported": suffix in SUPPORTED_CONTENT_EXTENSIONS,
            "skipped": False,
            "error": True,
            "reason": "extraction_error",
            "message": "Datei konnte nicht gelesen werden.",
            "extraction_status": "error",
            "text": "",
        }


def _extract_spreadsheet_with_metadata(path: Path) -> tuple[str, dict[str, int]]:
    """Extract once for metadata while maintaining the public string helper contract."""

    workbook = None
    try:
        workbook = load_workbook(path, read_only=True, data_only=True, keep_vba=False)
        return _spreadsheet_sections(workbook)
    finally:
        if workbook is not None:
            workbook.close()


def _extract_csv(path: Path) -> str:
    text, _metadata = _extract_csv_with_metadata(path)
    return text


def _extract_csv_with_metadata(path: Path) -> tuple[str, dict[str, int]]:
    content = _read_text_tolerant(path)
    sample = content[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel
    rows = list(csv.reader(io.StringIO(content), dialect))
    meaningful_rows = [row for row in rows if any(_cell_text(value) for value in row)]
    if not meaningful_rows:
        return "", {"row_count": 0}
    headers = _headers_from_row(meaningful_rows[0])
    data_rows = meaningful_rows[1:]
    return (
        "\n".join(
            _structured_row(row_number, headers, row)
            for row_number, row in enumerate(data_rows, start=1)
        ),
        {"row_count": len(data_rows)},
    )


def _spreadsheet_sections(workbook: Any) -> tuple[str, dict[str, int]]:
    sections: list[str] = []
    row_count = 0
    sheet_count = 0
    for sheet in workbook.worksheets:
        rows = [list(row) for row in sheet.iter_rows(values_only=True)]
        meaningful_rows = [row for row in rows if any(_has_value(value) for value in row)]
        if not meaningful_rows:
            continue
        sheet_count += 1
        headers = _headers_from_row(meaningful_rows[0])
        sheet_lines = [f"[Arbeitsblatt: {sheet.title}]"]
        if len(meaningful_rows) == 1:
            sheet_lines.append("Header | " + " | ".join(headers))
        for row_number, row in enumerate(meaningful_rows[1:], start=1):
            sheet_lines.append(_structured_row(row_number, headers, row))
            row_count += 1
        sections.append("\n".join(sheet_lines))
    return "\n\n".join(sections), {"sheet_count": sheet_count, "row_count": row_count}


def _read_text_tolerant(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def _bounded_json(data: Any) -> str:
    rendered = json.dumps(data, ensure_ascii=False, indent=2, default=str)
    return rendered[:_MAX_EXTRACTED_TEXT_CHARS]


def _headers_from_row(row: Iterable[Any]) -> list[str]:
    headers: list[str] = []
    used: set[str] = set()
    for position, value in enumerate(row, start=1):
        candidate = _cell_text(_format_value(value)) or f"Spalte {position}"
        base = candidate
        suffix = 2
        while candidate in used:
            candidate = f"{base} {suffix}"
            suffix += 1
        used.add(candidate)
        headers.append(candidate)
    return headers


def _structured_row(row_number: int, headers: list[str], row: Iterable[Any]) -> str:
    values = list(row)
    columns = []
    for position, value in enumerate(values):
        formatted = _cell_text(_format_value(value))
        if not formatted:
            continue
        header = headers[position] if position < len(headers) else f"Spalte {position + 1}"
        columns.append(f"{header}: {formatted}")
    return f"Zeile {row_number} | " + " | ".join(columns)


def _format_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "true" if value else "false"
    if isinstance(value, datetime):
        if value.time() == time.min:
            return value.date().isoformat()
        return value.isoformat(sep=" ", timespec="seconds")
    if isinstance(value, date):
        return value.isoformat()
    if isinstance(value, time):
        return value.isoformat(timespec="seconds")
    if isinstance(value, float):
        return f"{value:.2f}"
    return str(value)


def _has_value(value: Any) -> bool:
    return value is not None and (not isinstance(value, str) or value.strip() != "")


def _cell_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _is_heading(style_name: str) -> bool:
    normalized = style_name.casefold()
    return normalized.startswith("heading") or normalized.startswith("überschrift")


def _pdf_error(reason: str, message: str) -> dict[str, Any]:
    return {
        "success": False,
        "text": "",
        "error": True,
        "skipped": True,
        "reason": reason,
        "message": message,
        "page_count": 0,
        "extraction_status": "error",
    }


def clean_extracted_text(text: str) -> str:
    """Remove control characters while retaining paragraph boundaries for search snippets."""

    cleaned = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]+", " ", text)
    cleaned = re.sub(r"(\w)-[ \t]+(\w)", r"\1\2", cleaned)
    cleaned = re.sub(r"([a-zäöüß])([A-ZÄÖÜ])", r"\1 \2", cleaned)
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r" *\n *", "\n", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    return cleaned[:_MAX_EXTRACTED_TEXT_CHARS]


def clear_content_cache() -> None:
    _CONTENT_CACHE.clear()


def _content_cache_enabled() -> bool:
    return os.getenv("FILE_CONTENT_CACHE_ENABLED", "true").strip().lower() == "true"


def _content_cache_max_items() -> int:
    try:
        return max(0, int(os.getenv("FILE_CONTENT_CACHE_MAX_ITEMS", "200")))
    except ValueError:
        return 200


def _get_cached(key: tuple[str, int, int]) -> dict[str, Any] | None:
    if not _content_cache_enabled():
        return None
    cached = _CONTENT_CACHE.get(key)
    return dict(cached) if cached else None


def _set_cached(key: tuple[str, int, int], result: dict[str, Any]) -> None:
    if not _content_cache_enabled() or result.get("error") or result.get("skipped"):
        return
    max_items = _content_cache_max_items()
    if max_items <= 0:
        return
    while len(_CONTENT_CACHE) >= max_items:
        _CONTENT_CACHE.pop(next(iter(_CONTENT_CACHE)), None)
    # Cache only bounded, cleaned extraction results; never raw binaries or error details.
    _CONTENT_CACHE[key] = dict(result)
